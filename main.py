import os
import sys
import json
import tempfile
import cv2
import numpy as np
import pyautogui
import win32api
import win32con
from tkinter import *
from tkinter import filedialog
from PIL import Image, ImageTk
from docx.shared import Inches
import docx
import mouse
import keyboard
import threading
import time


# Import custom module to handle the display and edit of images and descriptions
import image_table_module
# Import custom module to handle loading created projects
import project_loader_module

screenshot_mode = False

# get information about the primary screen
primary_screen_info = win32api.GetMonitorInfo(win32api.EnumDisplayMonitors()[0][0])
primary_screen_left, primary_screen_top, primary_screen_right, primary_screen_bottom = primary_screen_info['Monitor']
primary_screen_width = primary_screen_right - primary_screen_left
primary_screen_height = primary_screen_bottom - primary_screen_top

# Function to get the highest number used in the filename
def get_highest_number():
    numbers = [int(filename.split(".")[0]) for filename in os.listdir(screenshots_folder) if filename.endswith(".png")]
    if numbers:
        return max(numbers)
    else:
        return 0

# Function to take a screenshot and mark the clicked spot
def take_screenshot(x, y):
    global num, screenshots_folder
    # get the current screen size and position
    screen_index = win32api.MonitorFromPoint((x, y), win32con.MONITOR_DEFAULTTONEAREST)
    screen_info = win32api.GetMonitorInfo(screen_index)
    screen_left, screen_top, screen_right, screen_bottom = screen_info['Monitor']
    screen_width = screen_right - screen_left
    screen_height = screen_bottom - screen_top
    # check if the clicked position is within the bounds of the primary screen
    if primary_screen_left <= x <= primary_screen_right and primary_screen_top <= y <= primary_screen_bottom:
        # take a screenshot of the screen that the mouse is currently on
        screenshot = pyautogui.screenshot(region=(screen_left, screen_top, screen_width, screen_height))
        # convert the screenshot to a numpy array
        screenshot = np.array(screenshot)
        # draw a red circle at the clicked spot
        cv2.circle(screenshot, (x - screen_left, y - screen_top), 10, (0, 0, 255), -1)
        # save the screenshot with the corresponding number
        filename = f"{num}.png"
        cv2.imwrite(os.path.join(screenshots_folder, filename), screenshot)
        print(f"Saved screenshot {os.path.join(screenshots_folder, filename)}")
        num += 1


def toggle_screenshot_mode():
    global screenshot_mode
    if not screenshot_mode:
        create_project_folder()
        screenshot_mode = True
        take_screenshot_button.config(text="Stop making screenshots")
        # Start a new thread to listen for mouse clicks
        t = threading.Thread(target=listen_for_mouse_clicks, daemon=True)
        t.start()
    else:
        screenshot_mode = False
        take_screenshot_button.config(text="Make screenshots")

def listen_for_mouse_clicks():
    while screenshot_mode:
        if win32api.GetAsyncKeyState(0x13) < 0:  # Stop the program if the pause/break key is pressed
            break
        if win32api.GetKeyState(0x01) < 0:  # Take a screenshot and mark the clicked spot if the left mouse button is pressed
            x, y = win32api.GetCursorPos()
            take_screenshot(x, y)
            time.sleep(0.1)

def stop_screenshot_mode():
    global screenshot_mode
    screenshot_mode = False
    cv2.destroyAllWindows()
    keyboard.unhook_all_hotkeys()

# Function to save the project as a Word document
def save_to_word():
    global project_name, descriptions
    doc = docx.Document()

    for image_file, description in descriptions.items():
        image_path = os.path.join(screenshots_folder, image_file)
        image = Image.open(image_path)
        doc.add_picture(image_path, width=Inches(6))
        paragraph = doc.add_paragraph()
        paragraph.add_run(description.strip())
        doc.add_page_break()

    if not os.path.exists("outputs"):
        os.makedirs("outputs")
    doc.save(os.path.join("outputs", f"{project_name}.docx"))

# Function to create project folder
def create_project_folder():
    global project_name, screenshots_folder, num
    project_name = project_name_var.get()
    if not project_name:
        return

    if not os.path.exists("projects"):
        os.makedirs("projects")

    project_folder = os.path.join("projects", project_name)
    if not os.path.exists(project_folder):
        os.makedirs(project_folder)

    screenshots_folder = os.path.join(project_folder, "screenshots")
    if not os.path.exists(screenshots_folder):
        os.makedirs(screenshots_folder)

    num = get_highest_number() + 1

# Function to load and display image table with descriptions
def load_image_table():
    global project_name, screenshots_folder, descriptions
    create_project_folder()
    descriptions_file = os.path.join(screenshots_folder, "descriptions.txt")
    json_file = os.path.join(screenshots_folder, "descriptions.json")

    image_files = sorted([filename for filename in os.listdir(screenshots_folder) if filename.endswith(".png")], key=lambda filename: int(filename.split(".")[0]))

    # Display the image table and edit descriptions using the custom module
    edited_descriptions = image_table_module.display_image_table(screenshots_folder, image_files, descriptions)

    # Update the descriptions if there are changes
    if edited_descriptions is not None:
        descriptions = edited_descriptions
        with open(json_file, "w") as f:
            json.dump(descriptions, f)

def open_project():
    selected_project = project_loader_module.load_project()

    if selected_project:
        project_name_var.set(selected_project)

# Create the main GUI window
root = Tk()
root.title("Auto instructions creator")

# Initialize variables
project_name_var = StringVar()
num = 1
screenshots_folder = ""
project_name = ""
descriptions = {}



# Create the project name entry and label
project_name_label = Label(root, text="Project name:")
project_name_label.grid(row=0, column=0, padx=5, pady=5)
project_name_entry = Entry(root, textvariable=project_name_var)
project_name_entry.grid(row=0, column=1, padx=5, pady=5)

# Add this line to the section where you create the buttons in the main script
open_project_button = Button(root, text="Open project", command=open_project)
open_project_button.grid(row=0, column=2, padx=5, pady=5)

# Create the buttons
take_screenshot_button = Button(root, text="Start making screenshots", command=toggle_screenshot_mode)
take_screenshot_button.grid(row=1, column=0, padx=5, pady=5)

load_image_table_button = Button(root, text="Edit descriptions", command=load_image_table)
load_image_table_button.grid(row=1, column=1, padx=5, pady=5)

save_to_word_button = Button(root, text="Save to Word", command=save_to_word)
save_to_word_button.grid(row=2, column=0, padx=5, pady=5)

# Run the main GUI loop
root.mainloop()